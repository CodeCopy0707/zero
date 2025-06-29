"""
Document processing tools for Agent Zero Gemini
"""
import asyncio
import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
import tempfile

from core.tools import BaseTool

logger = logging.getLogger(__name__)

class DocumentProcessorTool(BaseTool):
    """General document processing tool"""
    
    def __init__(self):
        super().__init__(
            name="document_processor",
            description="Process various document formats - extract text, metadata, convert formats"
        )
    
    async def execute(self, action: str, file_path: str, **kwargs) -> Dict[str, Any]:
        """Process document"""
        try:
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                return {
                    "success": False,
                    "error": f"File not found: {file_path}"
                }
            
            file_extension = file_path_obj.suffix.lower()
            
            if action == "extract_text":
                return await self._extract_text(file_path_obj, **kwargs)
            elif action == "get_metadata":
                return await self._get_metadata(file_path_obj, **kwargs)
            elif action == "convert":
                return await self._convert_document(file_path_obj, **kwargs)
            else:
                return {
                    "success": False,
                    "error": f"Unknown action: {action}"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _extract_text(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Extract text from document"""
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return await self._extract_pdf_text(file_path, **kwargs)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_word_text(file_path, **kwargs)
            elif file_extension in ['.xlsx', '.xls']:
                return await self._extract_excel_text(file_path, **kwargs)
            elif file_extension == '.txt':
                return await self._extract_plain_text(file_path, **kwargs)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file format: {file_extension}"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _extract_pdf_text(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            import PyPDF2
            
            text = ""
            page_texts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    page_texts.append({
                        "page": page_num + 1,
                        "text": page_text
                    })
                    text += page_text + "\n"
            
            return {
                "success": True,
                "text": text.strip(),
                "pages": page_texts,
                "page_count": len(pdf_reader.pages),
                "file": str(file_path)
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "PyPDF2 not installed. Install with: pip install PyPDF2"
            }
    
    async def _extract_word_text(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + "\n"
            
            return {
                "success": True,
                "text": full_text.strip(),
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "file": str(file_path)
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "python-docx not installed. Install with: pip install python-docx"
            }
    
    async def _extract_excel_text(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Extract text from Excel file"""
        try:
            import openpyxl
            
            workbook = openpyxl.load_workbook(file_path)
            
            sheets_data = {}
            all_text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = ""
                rows_data = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        rows_data.append(row_text)
                        sheet_text += row_text + "\n"
                
                sheets_data[sheet_name] = {
                    "text": sheet_text.strip(),
                    "rows": rows_data
                }
                all_text += f"Sheet: {sheet_name}\n{sheet_text}\n\n"
            
            return {
                "success": True,
                "text": all_text.strip(),
                "sheets": sheets_data,
                "sheet_count": len(workbook.sheetnames),
                "file": str(file_path)
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "openpyxl not installed. Install with: pip install openpyxl"
            }
    
    async def _extract_plain_text(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Extract text from plain text file"""
        encoding = kwargs.get("encoding", "utf-8")
        
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            lines = text.split('\n')
            
            return {
                "success": True,
                "text": text,
                "lines": lines,
                "line_count": len(lines),
                "file": str(file_path)
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=enc) as file:
                        text = file.read()
                    
                    return {
                        "success": True,
                        "text": text,
                        "encoding_used": enc,
                        "file": str(file_path)
                    }
                except UnicodeDecodeError:
                    continue
            
            return {
                "success": False,
                "error": "Could not decode file with any common encoding"
            }
    
    async def _get_metadata(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Get document metadata"""
        file_extension = file_path.suffix.lower()
        
        try:
            basic_metadata = {
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "file_extension": file_extension,
                "created": file_path.stat().st_ctime,
                "modified": file_path.stat().st_mtime
            }
            
            if file_extension == '.pdf':
                return await self._get_pdf_metadata(file_path, basic_metadata)
            elif file_extension in ['.docx', '.doc']:
                return await self._get_word_metadata(file_path, basic_metadata)
            else:
                return {
                    "success": True,
                    "metadata": basic_metadata
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _get_pdf_metadata(self, file_path: Path, basic_metadata: Dict) -> Dict[str, Any]:
        """Get PDF metadata"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = basic_metadata.copy()
                metadata.update({
                    "page_count": len(pdf_reader.pages),
                    "pdf_metadata": dict(pdf_reader.metadata) if pdf_reader.metadata else {}
                })
            
            return {
                "success": True,
                "metadata": metadata
            }
            
        except ImportError:
            return {
                "success": True,
                "metadata": basic_metadata,
                "warning": "PyPDF2 not available for detailed PDF metadata"
            }
    
    async def _get_word_metadata(self, file_path: Path, basic_metadata: Dict) -> Dict[str, Any]:
        """Get Word document metadata"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            metadata = basic_metadata.copy()
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "word_metadata": {
                    "author": doc.core_properties.author,
                    "title": doc.core_properties.title,
                    "subject": doc.core_properties.subject,
                    "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None
                }
            })
            
            return {
                "success": True,
                "metadata": metadata
            }
            
        except ImportError:
            return {
                "success": True,
                "metadata": basic_metadata,
                "warning": "python-docx not available for detailed Word metadata"
            }
    
    async def _convert_document(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Convert document format"""
        target_format = kwargs.get("target_format", "txt")
        output_file = kwargs.get("output_file")
        
        if not output_file:
            output_file = file_path.with_suffix(f".{target_format}")
        
        output_path = Path(output_file)
        
        try:
            # Extract text first
            text_result = await self._extract_text(file_path)
            
            if not text_result["success"]:
                return text_result
            
            text = text_result["text"]
            
            # Save in target format
            if target_format == "txt":
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            elif target_format == "html":
                html_content = f"<html><body><pre>{text}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported target format: {target_format}"
                }
            
            return {
                "success": True,
                "output_file": str(output_path),
                "target_format": target_format,
                "original_file": str(file_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def get_parameters(self) -> Dict[str, Any]:
        return {
            "action": {
                "type": "string",
                "description": "Document processing action",
                "enum": ["extract_text", "get_metadata", "convert"]
            },
            "file_path": {
                "type": "string",
                "description": "Path to document file"
            },
            "target_format": {
                "type": "string",
                "description": "Target format for conversion",
                "enum": ["txt", "html"],
                "optional": True
            },
            "output_file": {
                "type": "string",
                "description": "Output file path (optional)",
                "optional": True
            },
            "encoding": {
                "type": "string",
                "description": "Text encoding for plain text files",
                "default": "utf-8",
                "optional": True
            }
        }

class PDFTool(BaseTool):
    """Specialized PDF processing tool"""
    
    def __init__(self):
        super().__init__(
            name="pdf_processor",
            description="Advanced PDF processing - merge, split, extract pages, add watermarks"
        )
    
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """Process PDF"""
        try:
            import PyPDF2
            
            if action == "merge":
                return await self._merge_pdfs(**kwargs)
            elif action == "split":
                return await self._split_pdf(**kwargs)
            elif action == "extract_pages":
                return await self._extract_pages(**kwargs)
            elif action == "rotate_pages":
                return await self._rotate_pages(**kwargs)
            else:
                return {
                    "success": False,
                    "error": f"Unknown action: {action}"
                }
                
        except ImportError:
            return {
                "success": False,
                "error": "PyPDF2 not installed. Install with: pip install PyPDF2"
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _merge_pdfs(self, **kwargs) -> Dict[str, Any]:
        """Merge multiple PDF files"""
        import PyPDF2
        
        input_files = kwargs.get("input_files", [])
        output_file = kwargs.get("output_file", "merged.pdf")
        
        if not input_files:
            return {
                "success": False,
                "error": "No input files provided"
            }
        
        merger = PyPDF2.PdfMerger()
        
        try:
            for file_path in input_files:
                if Path(file_path).exists():
                    merger.append(file_path)
                else:
                    return {
                        "success": False,
                        "error": f"File not found: {file_path}"
                    }
            
            with open(output_file, 'wb') as output:
                merger.write(output)
            
            merger.close()
            
            return {
                "success": True,
                "output_file": output_file,
                "input_files": input_files,
                "message": f"Merged {len(input_files)} PDFs into {output_file}"
            }
            
        except Exception as e:
            merger.close()
            return {
                "success": False,
                "error": str(e)
            }
    
    def get_parameters(self) -> Dict[str, Any]:
        return {
            "action": {
                "type": "string",
                "description": "PDF processing action",
                "enum": ["merge", "split", "extract_pages", "rotate_pages"]
            },
            "input_files": {
                "type": "array",
                "description": "List of input PDF files (for merge action)",
                "optional": True
            },
            "output_file": {
                "type": "string",
                "description": "Output file path",
                "optional": True
            }
        }

class WordTool(BaseTool):
    """Specialized Word document processing tool"""
    
    def __init__(self):
        super().__init__(
            name="word_processor",
            description="Advanced Word document processing - create, modify, format documents"
        )
    
    async def execute(self, action: str, **kwargs) -> Dict[str, Any]:
        """Process Word document"""
        try:
            from docx import Document
            
            if action == "create":
                return await self._create_document(**kwargs)
            elif action == "add_content":
                return await self._add_content(**kwargs)
            elif action == "format_text":
                return await self._format_text(**kwargs)
            else:
                return {
                    "success": False,
                    "error": f"Unknown action: {action}"
                }
                
        except ImportError:
            return {
                "success": False,
                "error": "python-docx not installed. Install with: pip install python-docx"
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _create_document(self, **kwargs) -> Dict[str, Any]:
        """Create new Word document"""
        from docx import Document
        
        output_file = kwargs.get("output_file", "document.docx")
        title = kwargs.get("title", "")
        content = kwargs.get("content", "")
        
        doc = Document()
        
        if title:
            doc.add_heading(title, 0)
        
        if content:
            doc.add_paragraph(content)
        
        doc.save(output_file)
        
        return {
            "success": True,
            "output_file": output_file,
            "message": f"Created Word document: {output_file}"
        }
    
    def get_parameters(self) -> Dict[str, Any]:
        return {
            "action": {
                "type": "string",
                "description": "Word processing action",
                "enum": ["create", "add_content", "format_text"]
            },
            "output_file": {
                "type": "string",
                "description": "Output file path",
                "optional": True
            },
            "title": {
                "type": "string",
                "description": "Document title",
                "optional": True
            },
            "content": {
                "type": "string",
                "description": "Document content",
                "optional": True
            }
        }
